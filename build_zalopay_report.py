from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("output") / "Bao_cao_ca_nhan_ZaloPay_Sandbox.docx"

NAVY = "17365D"
BLUE = "2E74B5"
PALE_BLUE = "EAF2F8"
PALE_GRAY = "F2F4F7"
TEXT = "1F1F1F"
MUTED = "5B6573"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_cm, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_cm):
        grid_col.set(qn("w:w"), str(round(width / 2.54 * 1440)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(round(width / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, size=12, bold=False, italic=False, color=TEXT, font="Times New Roman"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc, text="", style=None, align=None, before=None, after=None, line=None, bold=False, italic=False, size=12, color=TEXT):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if line is not None:
        p.paragraph_format.line_spacing = line
    if text:
        set_run_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size={1: 15, 2: 13, 3: 12}[level], bold=True, color=NAVY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(text), size=12)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(text), size=12)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.right_indent = Cm(0.35)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F7FA")
    p_pr.append(shd)
    set_run_font(p.add_run(text), size=10, font="Consolas", color="283747")
    return p


def add_page_number(paragraph_obj):
    paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph_obj.add_run("Trang ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph_obj._p.append(field)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Báo cáo cá nhân - ZaloPay API Sandbox  |  ")
    set_run_font(run, size=9, bold=True, color=MUTED)
    run = p.add_run("Trường Đại học Khoa học Tự nhiên - ĐHQG HCM")
    set_run_font(run, size=9, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_table(doc, headers, rows, widths_cm, font_size=10.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_cm)
    head = table.rows[0]
    set_repeat_table_header(head)
    for cell, text in zip(head.cells, headers):
        set_cell_shading(cell, PALE_BLUE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=font_size, bold=True, color=NAVY)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run_font(p.add_run(text), size=font_size)
    paragraph(doc, "", after=1)
    return table


def set_document_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for level, size, before, after in ((1, 15, 15, 7), (2, 13, 11, 5), (3, 12, 8, 4)):
        style = styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(12)


def restart_page_numbering(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True
    set_document_styles(doc)
    add_header_footer(section)

    # Cover page
    for line in ["ĐẠI HỌC QUỐC GIA THÀNH PHỐ HỒ CHÍ MINH", "TRƯỜNG ĐẠI HỌC KHOA HỌC TỰ NHIÊN", "KHOA CÔNG NGHỆ THÔNG TIN", "BỘ MÔN HỆ THỐNG THÔNG TIN"]:
        paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, bold=True, size=12)
    paragraph(doc, "", after=42)
    paragraph(doc, "BÁO CÁO CÁ NHÂN", align=WD_ALIGN_PARAGRAPH.CENTER, after=12, bold=True, size=17, color=NAVY)
    paragraph(doc, "TÌM HIỂU VÀ THỰC HÀNH TÍCH HỢP\nZALOPAY API TRONG MÔI TRƯỜNG SANDBOX", align=WD_ALIGN_PARAGRAPH.CENTER, after=22, bold=True, size=16, color=NAVY, line=1.2)
    paragraph(doc, "Môn học: Thương mại điện tử", align=WD_ALIGN_PARAGRAPH.CENTER, after=36, italic=True, size=13)
    paragraph(doc, "Sinh viên thực hiện:", align=WD_ALIGN_PARAGRAPH.CENTER, after=3, bold=True, size=12)
    paragraph(doc, "Phạm Phát Lộc - 23127085", align=WD_ALIGN_PARAGRAPH.CENTER, after=18, size=13)
    paragraph(doc, "Giảng viên hướng dẫn:", align=WD_ALIGN_PARAGRAPH.CENTER, after=3, bold=True, size=12)
    paragraph(doc, "Lương Vĩ Minh và Nguyễn Đức Huy", align=WD_ALIGN_PARAGRAPH.CENTER, after=42, size=13)
    paragraph(doc, "Thành phố Hồ Chí Minh, tháng 8 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)
    doc.add_page_break()

    # Contents
    paragraph(doc, "MỤC LỤC", align=WD_ALIGN_PARAGRAPH.CENTER, after=14, bold=True, size=16, color=NAVY)
    toc_rows = [
        ("1. Giới thiệu", "1"), ("2. Tổng quan ZaloPay API", "1"), ("   2.1. Mô hình tích hợp thanh toán qua API", "1"),
        ("   2.2. Phân biệt Sandbox và Production", "1"), ("   2.3. Thông tin cấu hình", "1"),
        ("3. Quy trình thanh toán ZaloPay", "2"), ("4. API tạo đơn hàng", "2"),
        ("   4.1. Đặc tả request", "2"), ("   4.2. Mã tham chiếu giao dịch", "3"), ("   4.3. Tạo chữ ký MAC", "3"),
        ("5. Callback và bảo mật", "3"), ("6. QR động ZaloPay", "4"), ("7. Hướng dẫn demo tối thiểu", "4"),
        ("8. Đánh giá", "5"), ("9. Kết luận", "5"), ("Tài liệu tham khảo", "6"),
    ]
    toc = doc.add_table(rows=1, cols=2)
    toc.style = "Table Grid"
    set_table_geometry(toc, [14.3, 1.7])
    for c, t in zip(toc.rows[0].cells, ("Nội dung", "Trang")):
        set_cell_shading(c, PALE_BLUE)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(c.paragraphs[0].add_run(t), size=11, bold=True, color=NAVY)
    for title, page in toc_rows:
        cells = toc.add_row().cells
        set_run_font(cells[0].paragraphs[0].add_run(title), size=10.5)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cells[1].paragraphs[0].add_run(page), size=10.5)
    content_section = doc.add_section(WD_SECTION.NEW_PAGE)
    content_section.page_width = Cm(21)
    content_section.page_height = Cm(29.7)
    content_section.top_margin = Cm(2.5)
    content_section.bottom_margin = Cm(2.3)
    content_section.left_margin = Cm(2.5)
    content_section.right_margin = Cm(2.5)
    content_section.header_distance = Cm(1.25)
    content_section.footer_distance = Cm(1.25)
    restart_page_numbering(content_section, 1)
    add_heading(doc, "1. Giới thiệu")
    paragraph(doc, "ZaloPay là nền tảng thanh toán điện tử tại Việt Nam, hỗ trợ thanh toán trên ứng dụng, website và điểm bán. Với lập trình viên, ZaloPay cung cấp API để hệ thống tạo yêu cầu thanh toán, chuyển người dùng vào luồng thanh toán và tiếp nhận kết quả giao dịch từ phía máy chủ.")
    paragraph(doc, "Báo cáo này tìm hiểu quy trình tích hợp ZaloPay API ở mức tối thiểu nhưng đầy đủ về mặt kỹ thuật: tạo giao dịch trên backend, nhận order_url, hiển thị trang thanh toán hoặc QR động, tiếp nhận callback và xác thực kết quả bằng HMAC-SHA256. Trọng tâm là chỉ ghi nhận đơn hàng khi dữ liệu phía máy chủ đã được xác minh.")
    paragraph(doc, "Toàn bộ nội dung thực hành sử dụng ZaloPay Sandbox, là môi trường dành cho học tập và kiểm thử với dữ liệu mô phỏng, không phát sinh giao dịch tiền thật.")

    add_heading(doc, "2. Tổng quan ZaloPay API")
    add_heading(doc, "2.1. Mô hình tích hợp thanh toán qua API", 2)
    paragraph(doc, "Trong mô hình tích hợp qua API, ứng dụng không trực tiếp xử lý thông tin thanh toán nhạy cảm. Backend chuẩn bị đơn hàng, ký request bằng khóa bí mật và gửi yêu cầu tới ZaloPay. Khi request hợp lệ, ZaloPay trả về dữ liệu để người dùng tiếp tục thanh toán; sau đó hệ thống ZaloPay gọi callback đến backend để thông báo kết quả.")
    paragraph(doc, "Mô hình tạo thành hai luồng riêng: luồng trình duyệt phục vụ trải nghiệm người dùng qua order_url và Redirect URL; luồng server-to-server dùng callback để xác nhận kết quả. Backend cần lưu mã giao dịch nội bộ để liên kết hai luồng và tránh cập nhật nhầm đơn.")

    add_heading(doc, "2.2. Phân biệt Sandbox và Production", 2)
    add_table(doc, ["Môi trường", "Mục đích", "Đặc điểm sử dụng"], [
        ("Sandbox", "Học tập, phát triển và kiểm thử luồng tích hợp.", "Dùng App ID và khóa Sandbox; dữ liệu thanh toán chỉ mang tính thử nghiệm."),
        ("Production", "Tiếp nhận giao dịch thực tế.", "Dành cho merchant/doanh nghiệp đã hoàn tất đăng ký, kiểm tra và được cấp thông tin chính thức."),
    ], [3.0, 5.3, 7.7], 10.5)
    paragraph(doc, "Endpoint tạo đơn minh họa trong báo cáo: POST https://sb-openapi.zalopay.vn/v2/create", italic=True, size=10.5, color=MUTED)

    add_heading(doc, "2.3. Thông tin cấu hình", 2)
    add_table(doc, ["Biến cấu hình", "Vai trò"], [
        ("ZALOPAY_APP_ID", "Định danh ứng dụng Sandbox hoặc merchant do ZaloPay cấp."),
        ("ZALOPAY_KEY1", "Khóa tạo chữ ký MAC cho request gửi tới ZaloPay."),
        ("ZALOPAY_KEY2", "Khóa kiểm tra chữ ký MAC của callback nhận từ ZaloPay."),
        ("ZALOPAY_API_URL", "Địa chỉ API tạo đơn của môi trường đang sử dụng."),
        ("ZALOPAY_CALLBACK_URL", "Endpoint backend công khai để nhận kết quả giao dịch."),
        ("ZALOPAY_REDIRECT_URL", "Trang người dùng quay lại sau khi hoàn tất hoặc hủy thanh toán."),
    ], [5.4, 10.6])
    paragraph(doc, "Key1 và Key2 là thông tin bảo mật phía máy chủ; không được đưa vào frontend, kho mã nguồn công khai, ảnh minh họa hoặc phụ lục. Báo cáo không sử dụng bất kỳ giá trị khóa thật nào.")

    add_heading(doc, "3. Quy trình thanh toán ZaloPay")
    paragraph(doc, "Luồng thanh toán tối thiểu được mô tả như sau:", after=4)
    for step in ["Người dùng khởi tạo thanh toán", "Backend tạo đơn và ký request bằng Key1", "Gửi request tới API ZaloPay Sandbox /v2/create", "Nhận order_url; mở trang thanh toán hoặc tạo QR động", "ZaloPay gửi callback tới backend", "Backend kiểm tra MAC bằng Key2 và đối chiếu dữ liệu", "Lưu trạng thái thành công/thất bại, sau đó điều hướng người dùng về Redirect URL"]:
        paragraph(doc, step, align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1, bold=True, size=11, color=NAVY)
        if step != "Lưu trạng thái thành công/thất bại, sau đó điều hướng người dùng về Redirect URL":
            paragraph(doc, "↓", align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, bold=True, size=12, color=BLUE)
    paragraph(doc, "Hình 1. Luồng thanh toán ZaloPay Sandbox trong demo tối thiểu", align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=7, italic=True, size=10, color=MUTED)
    paragraph(doc, "Redirect URL chỉ phục vụ điều hướng và hiển thị kết quả. Người dùng có thể tự mở lại trang này khi giao dịch chưa hoàn tất, vì vậy đây không phải là bằng chứng thanh toán đáng tin cậy. Callback đã xác thực MAC và được đối chiếu với dữ liệu nội bộ mới là căn cứ xác nhận giao dịch.")

    add_heading(doc, "4. API tạo đơn hàng")
    add_heading(doc, "4.1. Đặc tả request", 2)
    paragraph(doc, "API tạo đơn dùng phương thức POST tại path /v2/create. Request minh họa được gửi dưới định dạng application/x-www-form-urlencoded.")
    add_table(doc, ["Tham số", "Ý nghĩa"], [
        ("app_id", "Mã ứng dụng ZaloPay được cấp cho môi trường tích hợp."),
        ("app_trans_id", "Mã giao dịch do hệ thống tự tạo, có tiền tố ngày hiện tại và không trùng trong ứng dụng."),
        ("app_user", "Mã hoặc thông tin định danh người dùng thực hiện giao dịch."),
        ("app_time", "Thời điểm tạo đơn dưới dạng Unix timestamp, tính bằng millisecond."),
        ("amount", "Số tiền thanh toán, đơn vị VND."),
        ("item", "Danh sách sản phẩm hoặc dịch vụ dưới dạng chuỗi JSON."),
        ("embed_data", "Dữ liệu bổ sung dạng JSON; có thể chứa redirecturl."),
        ("description", "Nội dung mô tả đơn hàng hiển thị trong luồng thanh toán."),
        ("bank_code", "Kênh thanh toán; có thể để trống tùy mô hình."),
        ("mac", "Chữ ký HMAC-SHA256 bảo vệ tính toàn vẹn của request."),
    ], [4.5, 11.5], 10.2)
    add_heading(doc, "4.2. Mã tham chiếu giao dịch", 2)
    paragraph(doc, "app_trans_id có quy ước dạng YYMMDD_xxx. Ví dụ:")
    add_code(doc, "260828_demo001")
    paragraph(doc, "Trong ví dụ này, 260828 biểu diễn ngày 28/08/2026; demo001 là phần mã riêng giúp tránh trùng lặp. Trong ứng dụng thực tế, phần mã riêng nên được sinh bằng ID đủ mạnh hoặc lấy từ khóa chính duy nhất thay vì tăng thủ công.")
    add_heading(doc, "4.3. Tạo chữ ký MAC", 2)
    paragraph(doc, "Dữ liệu đầu vào phải ghép đúng thứ tự và giữ nguyên nội dung của embed_data, item:")
    add_code(doc, "data = app_id|app_trans_id|app_user|amount|app_time|embed_data|item")
    add_code(doc, "mac = HMAC_SHA256(Key1, data)")
    paragraph(doc, "Key1 chỉ nên được đọc từ biến môi trường khi chạy backend. Đặt khóa trực tiếp trong mã nguồn làm tăng nguy cơ lộ secret qua lịch sử phiên bản, log hoặc gói frontend.")

    add_heading(doc, "5. Callback và bảo mật")
    add_heading(doc, "5.1. Cấu trúc callback", 2)
    paragraph(doc, "Sau khi giao dịch hoàn tất, ZaloPay gửi callback đến endpoint mà backend cung cấp. Callback có Content-Type: application/json và thường có ba trường bao ngoài:")
    add_table(doc, ["Trường", "Ý nghĩa"], [
        ("data", "Chuỗi JSON chứa dữ liệu giao dịch được callback."),
        ("mac", "Chữ ký dùng để xác minh tính hợp lệ của trường data."),
        ("type", "Loại callback; với callback đơn hàng, thường có giá trị 1."),
    ], [3.5, 12.5])
    paragraph(doc, "Khi phân tích data, backend có thể nhận các trường như app_trans_id, amount, zp_trans_id, app_time, server_time và channel. Không nên ghi toàn bộ payload vào log khi có dữ liệu không cần thiết.")
    add_heading(doc, "5.2. Xác thực callback", 2)
    paragraph(doc, "Backend cần tính HMAC trực tiếp trên chuỗi data nhận được, không phải trên đối tượng JSON đã parse rồi serialize lại:")
    add_code(doc, "expected_mac = HMAC_SHA256(Key2, data)")
    paragraph(doc, "Nếu MAC sai, backend phải từ chối callback và không cập nhật trạng thái giao dịch. Nếu MAC đúng, backend tiếp tục tìm giao dịch theo app_trans_id, kiểm tra số tiền đã lưu, zp_trans_id và trạng thái hiện tại. Xử lý nên có tính idempotent để callback lặp không tạo thêm tác vụ nghiệp vụ.")

    add_heading(doc, "6. QR động ZaloPay")
    paragraph(doc, "QR động trong luồng này không phải QR nhận tiền cá nhân. Sau khi API tạo đơn thành công, ZaloPay trả về order_url; ứng dụng có thể dùng thư viện QR để mã hóa nguyên văn URL này. Khi người dùng quét mã, ứng dụng thanh toán sẽ mở đúng luồng giao dịch đã được tạo.")
    paragraph(doc, "Mỗi QR động gắn với một yêu cầu thanh toán cụ thể qua order_url, số tiền và app_trans_id. QR Sandbox chỉ phục vụ giao dịch thử nghiệm, không thay thế QR tĩnh nhận tiền tại quầy của merchant.")
    for item in ["Backend tạo đơn Sandbox và nhận order_url.", "Frontend/trang demo mã hóa nguyên văn order_url thành QR.", "Người dùng quét QR và thực hiện giao dịch thử nghiệm.", "Backend nhận callback, xác thực MAC rồi mới hiển thị trạng thái hoàn tất."]:
        add_number(doc, item)

    add_heading(doc, "7. Hướng dẫn demo tối thiểu")
    add_heading(doc, "7.1. Clip demo", 2)
    paragraph(doc, "Tên clip: 23127085-Zalopay. Liên kết YouTube: https://youtu.be/mzfwhj3oZZY")
    add_heading(doc, "7.2. Bước 1: Chuẩn bị thông tin Sandbox", 2)
    paragraph(doc, "Đăng ký hoặc nhận App ID, Key1 và Key2 theo hướng dẫn Sandbox. Lưu thông tin trong biến môi trường của backend, giới hạn quyền truy cập và kiểm tra đúng cặp App ID/khóa của cùng môi trường.")
    add_heading(doc, "7.3. Bước 2: Cấu hình biến môi trường", 2)
    add_code(doc, 'ZALOPAY_APP_ID="YOUR_SANDBOX_APP_ID"\nZALOPAY_KEY1="YOUR_KEY1"\nZALOPAY_KEY2="YOUR_KEY2"\nZALOPAY_API_URL="https://sb-openapi.zalopay.vn/v2/create"\nZALOPAY_CALLBACK_URL="https://your-public-domain/payments/zalopay/callback"\nZALOPAY_REDIRECT_URL="http://localhost:3000/payments/return/zalopay"')
    paragraph(doc, "Callback URL phải là endpoint backend công khai để ZaloPay gửi POST từ bên ngoài. Redirect URL có thể là localhost trong lúc demo giao diện, nhưng localhost không phù hợp cho callback.")
    add_heading(doc, "7.4. Bước 3: Tạo đơn hàng", 2)
    paragraph(doc, "Backend sinh app_trans_id duy nhất, chuẩn bị amount, item, embed_data và description. Sau đó ghép chuỗi đúng thứ tự, tạo MAC bằng Key1 và gửi POST form-urlencoded. Khi return_code = 1, lấy order_url từ phản hồi để tiếp tục.")
    add_heading(doc, "7.5. Bước 4: Hiển thị và thanh toán", 2)
    paragraph(doc, "Trang demo mở trực tiếp order_url hoặc tạo QR động từ URL này. Trong khi chưa nhận callback, giao diện nên hiển thị trạng thái đang chờ.")
    add_heading(doc, "7.6. Bước 5: Nhận callback", 2)
    paragraph(doc, "Backend nhận data và mac, tính expected_mac bằng Key2 và so sánh an toàn. Nếu hợp lệ, đối chiếu mã tham chiếu, số tiền, lưu trạng thái và zp_trans_id. Nếu không hợp lệ, trả phản hồi thất bại và giữ nguyên trạng thái cũ.")
    add_heading(doc, "7.7. Bước 6: Kiểm tra kết quả", 2)
    add_table(doc, ["Điểm kiểm tra", "Kết quả mong đợi", "Xử lý khi không đạt"], [
        ("Tạo đơn Sandbox", "return_code = 1 và có order_url.", "Kiểm tra App ID, endpoint, dữ liệu MAC và định dạng request."),
        ("Mã tham chiếu", "Đúng dạng ngày và không trùng.", "Thay cơ chế sinh mã riêng; kiểm tra múi giờ Việt Nam."),
        ("Callback hợp lệ", "MAC đúng, mã và số tiền khớp.", "Kiểm tra Key2, chuỗi data nguyên văn và callback URL."),
        ("Callback giả mạo", "Bị từ chối, trạng thái không đổi.", "Không bỏ qua so sánh MAC và đối chiếu dữ liệu nội bộ."),
        ("Callback lặp", "Chỉ ghi nhận một lần.", "Bổ sung kiểm tra trạng thái và ràng buộc duy nhất cho giao dịch."),
    ], [4.0, 5.3, 5.7], 9.7)
    paragraph(doc, "Bảng 1. Danh sách kiểm tra cho demo ZaloPay Sandbox", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10, color=MUTED)

    add_heading(doc, "8. Đánh giá")
    add_heading(doc, "8.1. Ưu điểm", 2)
    for item in ["Hỗ trợ luồng trang thanh toán và QR động; order_url có thể dùng cho cả chuyển tiếp và tạo QR.", "Sandbox giúp kiểm thử quy trình trước khi tiếp cận Production.", "Callback server-to-server đáng tin cậy hơn redirect phía trình duyệt.", "HMAC-SHA256 cùng Key1/Key2 bảo vệ tính toàn vẹn và nguồn gốc dữ liệu."]:
        add_bullet(doc, item)
    add_heading(doc, "8.2. Hạn chế", 2)
    for item in ["Callback URL cần truy cập công khai; khi triển khai nên dùng HTTPS và domain ổn định.", "Khóa bí mật cần được quản lý kỹ trong phát triển, triển khai và ghi log.", "Sandbox không thay thế toàn bộ quy trình đăng ký, kiểm tra và vận hành Production.", "QR Sandbox chỉ phục vụ giao dịch thử nghiệm, không phải QR nhận tiền vào ví cá nhân."]:
        add_bullet(doc, item)
    add_heading(doc, "8.3. Bài học rút ra", 2)
    paragraph(doc, "Điểm quan trọng nhất là thiết kế mã giao dịch duy nhất ngay từ đầu và dùng mã này để đối chiếu mọi thông báo. Backend cần xác thực callback, kiểm tra số tiền và ngăn xử lý lặp trước khi ghi nhận thành công. Secret phải được tách khỏi frontend và mã nguồn; trạng thái hiển thị cho người dùng phải phản ánh dữ liệu đã được backend xác minh.")

    add_heading(doc, "9. Kết luận")
    paragraph(doc, "Báo cáo đã trình bày quy trình tích hợp ZaloPay API trong Sandbox từ chuẩn bị cấu hình, tạo đơn /v2/create, tạo MAC bằng Key1, xử lý order_url/QR động đến tiếp nhận và xác thực callback bằng Key2. Một tích hợp thanh toán đúng không chỉ mở được giao diện thanh toán mà còn phải xác minh kết quả bằng dữ liệu server-to-server.")
    paragraph(doc, "ZaloPay Sandbox phù hợp để thực hành tạo mã tham chiếu, ký HMAC-SHA256, quản lý trạng thái và kiểm tra callback mà không phát sinh tiền thật. Hướng phát triển tiếp theo là hoàn tất quy trình merchant Production, bổ sung API tra cứu trạng thái khi callback đến trễ và nghiên cứu luồng hoàn tiền theo đặc tả chính thức.")

    add_heading(doc, "Tài liệu tham khảo")
    references = [
        "[1] ZaloPay Developer Documentation, Tổng quan tích hợp ZaloPay. https://developers.zalopay.vn/v2/general/overview.html",
        "[2] ZaloPay Documentation, Tạo đơn hàng. https://docs.zalopay.vn/docs/specs/order-create/",
        "[3] ZaloPay Documentation, Callback. https://docs.zalopay.vn/vi/docs/developer-tools/knowledge-base/callback/",
        "[4] ZaloPay Documentation, Bảo mật truyền dữ liệu. https://docs.zalopay.vn/vi/docs/developer-tools/security/secure-data-transmission/",
        "[5] ZaloPay Documentation, Mã QR động. https://docs.zalopay.vn/vi/docs/guides/payment-acceptance/qrcode/dynamic-qr/",
        "[6] ZaloPay Developer Documentation, Hướng dẫn QR tĩnh. https://developers.zalopay.vn/v2/docs/staticqr/tutorial.html",
    ]
    for ref in references:
        paragraph(doc, ref, after=4, size=10.5)

    doc.core_properties.title = "Báo cáo cá nhân - ZaloPay API Sandbox"
    doc.core_properties.author = "Phạm Phát Lộc"
    doc.core_properties.subject = "Thương mại điện tử"
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
